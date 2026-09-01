import base64
import binascii
import io
import re
from typing import Any

import pymupdf
import spacy
from docx import Document
from fastapi import FastAPI, HTTPException
from pydantic import BaseModel, Field
from spacy.matcher import PhraseMatcher

APP_VERSION = "1.0.0"
MAX_FILE_BYTES = 3 * 1024 * 1024
MAX_TEXT_CHARS = 50_000

SECTION_DEFINITIONS = [
    ("summary", "Professional summary", ["summary", "professional summary", "career summary", "profile", "professional profile", "objective", "career objective"]),
    ("skills", "Technical skills", ["skills", "technical skills", "skills and technologies", "skills and tools", "technologies", "core competencies"]),
    ("experience", "Experience or projects", ["experience", "professional experience", "work experience", "work history", "employment", "internship", "projects", "project experience"]),
    ("education", "Education", ["education", "education and qualifications", "academic", "academic background", "qualifications"]),
    ("certifications", "Certifications", ["certifications", "certification", "certifications and achievements", "achievements", "courses and certifications"]),
]

SKILL_CATALOG = {
    "JavaScript": ["javascript"], "TypeScript": ["typescript"], "React": ["react", "react.js"], "Node.js": ["node.js", "nodejs"],
    "Express.js": ["express", "express.js"], "MongoDB": ["mongodb"], "SQL": ["sql", "mysql", "postgresql"], "Python": ["python"],
    "Java": ["java"], "C++": ["c++", "cpp"], "REST APIs": ["rest api", "restful api", "restful"], "Git": ["git", "github"],
    "Docker": ["docker"], "AWS": ["aws", "amazon web services"], "Testing": ["jest", "testing", "unit test", "pytest"],
    "Data Structures & Algorithms": ["data structures", "algorithms", "dsa"], "System Design": ["system design"],
    "Next.js": ["next.js", "nextjs"], "Angular": ["angular"], "Vue.js": ["vue.js", "vuejs"], "HTML/CSS": ["html", "css"],
    "Redux": ["redux"], "GraphQL": ["graphql"], "Postman": ["postman"], "Linux": ["linux"],
    "Kubernetes": ["kubernetes", "k8s"], "CI/CD": ["ci/cd", "continuous integration", "continuous deployment"],
    "Azure": ["azure"], "GCP": ["google cloud", "gcp"], "Figma": ["figma"], "Agile": ["agile", "scrum"],
    "Machine Learning": ["machine learning", "ml"], "Data Analysis": ["data analysis", "pandas", "power bi"],
}
ACTION_VERBS = {"built", "developed", "designed", "implemented", "engineered", "optimized", "improved", "created", "led", "delivered", "integrated", "automated", "deployed", "reduced", "increased", "architected", "spearheaded"}

nlp = spacy.blank("en")
skill_matcher = PhraseMatcher(nlp.vocab, attr="LOWER")
for skill, aliases in SKILL_CATALOG.items():
    skill_matcher.add(skill, [nlp.make_doc(alias) for alias in aliases])


class ResumeAnalysisRequest(BaseModel):
    file_base64: str = Field(min_length=1, max_length=4_300_000)
    file_name: str = Field(min_length=1, max_length=255)
    mime_type: str = Field(min_length=1, max_length=120)


app = FastAPI(title="CareerPilot Resume Analysis Service", version=APP_VERSION)


def clamp(value: float) -> int:
    return max(0, min(100, round(value)))


def normalize(value: str) -> str:
    return re.sub(r"\s+", " ", value.replace("\r", " ")).strip()


def lines_from(text: str) -> list[str]:
    return [line.strip() for line in text.replace("\r", "").split("\n") if line.strip()]


def is_pdf(name: str, mime_type: str) -> bool:
    return mime_type == "application/pdf" or name.lower().endswith(".pdf")


def parse_pdf(buffer: bytes) -> tuple[str, dict[str, Any]]:
    document = pymupdf.open(stream=buffer, filetype="pdf")
    chunks: list[str] = []
    warnings: list[str] = []
    image_count = 0
    block_count = 0
    multi_column_pages = 0
    page_count = len(document)
    try:
        for page_index, page in enumerate(document, start=1):
            blocks = [block for block in page.get_text("blocks", sort=True) if str(block[4]).strip()]
            chunks.extend(str(block[4]).strip() for block in blocks)
            block_count += len(blocks)
            image_count += len(page.get_images(full=True))
            midpoint = page.rect.width / 2
            left_blocks = sum(1 for block in blocks if block[0] < midpoint * 0.92)
            right_blocks = sum(1 for block in blocks if block[2] > midpoint * 1.08)
            if left_blocks >= 3 and right_blocks >= 3:
                multi_column_pages += 1
                warnings.append(f"Potential multi-column reading order detected on page {page_index}.")
    finally:
        document.close()
    return "\n".join(chunks)[:MAX_TEXT_CHARS], {
        "pageCount": page_count, "blockCount": block_count, "imageCount": image_count,
        "tableCount": 0, "multiColumnPages": multi_column_pages, "layoutWarnings": warnings[:4],
    }


def parse_docx(buffer: bytes) -> tuple[str, dict[str, Any]]:
    document = Document(io.BytesIO(buffer))
    parts = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    table_count = len(document.tables)
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    warnings = ["Tables were detected; use simple text sections when ATS compatibility is critical."] if table_count else []
    return "\n".join(parts)[:MAX_TEXT_CHARS], {
        "pageCount": 0, "blockCount": len(parts), "imageCount": 0,
        "tableCount": table_count, "multiColumnPages": 0, "layoutWarnings": warnings,
    }


def extract_document(buffer: bytes, file_name: str, mime_type: str) -> tuple[str, dict[str, Any]]:
    if is_pdf(file_name, mime_type):
        return parse_pdf(buffer)
    if mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document" or file_name.lower().endswith(".docx"):
        return parse_docx(buffer)
    raise ValueError("Only PDF and DOCX files are supported.")


def detected_sections(lines: list[str]) -> list[dict[str, Any]]:
    normalized_lines = [re.sub(r"[^a-z0-9 ]", " ", line.lower()).strip() for line in lines]
    results = []
    for key, label, aliases in SECTION_DEFINITIONS:
        present = any(len(line) <= 48 and any(line == alias or line.startswith(f"{alias} ") or line.endswith(f" {alias}") for alias in aliases) for line in normalized_lines)
        results.append({"key": key, "label": label, "present": present})
    return results


def extract_skills(text: str) -> list[str]:
    document = nlp(text)
    return sorted({nlp.vocab.strings[match_id] for match_id, _, _ in skill_matcher(document)})


def build_analysis(text: str, document_data: dict[str, Any]) -> dict[str, Any]:
    normalized = normalize(text)
    lines = lines_from(text)
    word_count = len(normalized.split())
    sections = detected_sections(lines)
    core_sections = sections[:4]
    skills = extract_skills(normalized)
    bullet_lines = [line for line in lines if re.match(r"^(?:[-*\u2022\u25aa\u25e6]|\d+[.)])\s+", line)]
    action_led = [line for line in bullet_lines if re.sub(r"^(?:[-*\u2022\u25aa\u25e6]|\d+[.)])\s+", "", line).split(" ", 1)[0].lower() in ACTION_VERBS]
    metric_count = len(re.findall(r"\b\d+(?:\.\d+)?\s?(?:%|x|ms|seconds?|minutes?|hours?|users?|requests?|projects?|days?|months?|years?)\b", normalized, flags=re.I))
    timeline_count = len(re.findall(r"\b(?:(?:jan(?:uary)?|feb(?:ruary)?|mar(?:ch)?|apr(?:il)?|may|jun(?:e)?|jul(?:y)?|aug(?:ust)?|sep(?:t(?:ember)?)?|oct(?:ober)?|nov(?:ember)?|dec(?:ember)?)[\s,]*)?(?:19|20)\d{2}\s*(?:-|to)\s*(?:(?:jan(?:uary)?|feb(?:ruary)?|mar(?:ch)?|apr(?:il)?|may|jun(?:e)?|jul(?:y)?|aug(?:ust)?|sep(?:t(?:ember)?)?|oct(?:ober)?|nov(?:ember)?|dec(?:ember)?)[\s,]*)?(?:present|current|(?:19|20)\d{2})\b", normalized, flags=re.I))
    text_ratio = len(re.findall(r"[\w\s.,;:()/%+&@#'\-]", normalized)) / max(len(normalized), 1)
    parser_score = 0 if word_count < 80 else clamp(min(word_count / 250, 1) * 45 + text_ratio * 35 + (20 if not document_data["layoutWarnings"] else 8))
    layout_score = clamp(100 - document_data["multiColumnPages"] * 26 - document_data["tableCount"] * 12 - (12 if document_data["imageCount"] > 10 else 0))
    section_score = clamp((sum(section["present"] for section in core_sections) / len(core_sections)) * 80 + (20 if sections[-1]["present"] else 0))
    evidence_score = clamp(min(len(action_led) / 3, 1) * 42 + min(metric_count / 3, 1) * 38 + min(len(bullet_lines) / 4, 1) * 20)
    skills_score = clamp(min(len(skills) / 8, 1) * 70 + (30 if sections[1]["present"] else 0))
    chronology_score = clamp(min(timeline_count / 2, 1) * 100)
    overall_score = clamp(parser_score * .25 + section_score * .20 + chronology_score * .10 + layout_score * .15 + skills_score * .10 + evidence_score * .20)
    findings: list[dict[str, Any]] = []

    def add_finding(identifier: str, category: str, priority: str, title: str, detail: str, impact: int, evidence: str | None = None) -> None:
        findings.append({"id": identifier, "category": category, "priority": priority, "title": title, "detail": detail, "scoreImpact": impact, "evidence": evidence})

    if word_count < 80:
        add_finding("python-parser-text", "Parser health", "critical", "Resume text is not reliably extractable", "Very little readable text was extracted. Use a text-based PDF or a standard DOCX file.", 20, f"{word_count} words extracted")
    for index, warning in enumerate(document_data["layoutWarnings"]):
        add_finding(f"python-layout-{index}", "Document layout", "high", "Simplify the document layout", "Complex columns or tables can change reading order in applicant-tracking systems. Use a single-column, text-first layout for the safest parsing.", 10, warning)
    if len(action_led) < 2:
        add_finding("python-action-bullets", "Achievement quality", "medium", "Lead achievements with specific actions", "Start experience and project bullets with clear action verbs, then state the scope and result.", 5, f"{len(action_led)} action-led bullet lines detected")
    if metric_count < 2:
        add_finding("python-metrics", "Achievement quality", "medium", "Quantify outcomes where evidence exists", "Add genuine scope, time, performance, adoption, or delivery metrics to the accomplishments that can be measured.", 7, f"{metric_count} measurable outcomes detected")
    if timeline_count < 1:
        add_finding("python-chronology", "Timeline clarity", "low", "Use clear date ranges", "Add consistent month/year or year ranges for education, work, internships, and substantial projects.", 3, "No clear date range detected")
    if len(skills) < 4:
        add_finding("python-skills", "Skills organization", "medium", "Make skills easier to verify", "Use a dedicated Technical Skills section and group tools by category so they are easy to scan.", 5, f"{len(skills)} recognized skills detected")
    return {
        # Preserve line breaks so the Node fallback scorer can continue to
        # identify conventional section headings from the richer extraction.
        "engineVersion": APP_VERSION, "text": text.strip()[:MAX_TEXT_CHARS],
        "scores": {"parseability": parser_score, "layout": layout_score, "sections": section_score, "evidence": evidence_score, "skills": skills_score, "chronology": chronology_score, "overall": overall_score},
        "skills": skills, "sections": sections, "findings": findings[:12], "document": document_data,
    }


@app.get("/health")
def health() -> dict[str, str]:
    return {"status": "ok", "version": APP_VERSION}


@app.post("/analyze")
def analyze(request: ResumeAnalysisRequest) -> dict[str, Any]:
    try:
        buffer = base64.b64decode(request.file_base64, validate=True)
    except (binascii.Error, ValueError) as error:
        raise HTTPException(status_code=400, detail="Invalid resume payload.") from error
    if len(buffer) > MAX_FILE_BYTES:
        raise HTTPException(status_code=413, detail="Resume exceeds the 3 MB limit.")
    try:
        text, document_data = extract_document(buffer, request.file_name, request.mime_type)
        return build_analysis(text, document_data)
    except ValueError as error:
        raise HTTPException(status_code=400, detail=str(error)) from error
    except Exception as error:
        raise HTTPException(status_code=422, detail="The resume could not be analyzed by the document service.") from error
